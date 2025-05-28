import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

EXPORT_DIR = "static/reports"
os.makedirs(EXPORT_DIR, exist_ok=True)

def export_report_to_word(text: str, filename: str) -> str:
    path = os.path.join(EXPORT_DIR, f"{filename}.docx")
    doc = Document()
    doc.add_heading("运营分析报告", level=1)
    for para in text.split("\n"):
        doc.add_paragraph(para)
    doc.save(path)
    return path

def export_report_to_pdf(text: str, filename: str) -> str:
    path = os.path.join(EXPORT_DIR, f"{filename}.pdf")
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    y = height - 40
    for line in text.split("\n"):
        c.drawString(40, y, line)
        y -= 14
    c.save()
    return path