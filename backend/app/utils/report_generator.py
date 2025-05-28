from docx import Document
from docx.shared import Inches
from datetime import datetime
import os


def generate_word_report_with_chart(
    feedback_text: str,
    social_media_text: str,
    llm_summary: str,
    chart_path: str = None,
    filename: str = "report.docx"
) -> str:
    """
    生成包含封面、目录、客户反馈、社交媒体分析、图表（可选）、LLM总结的Word报告
    """
    doc = Document()

    # 1. 封面页
    doc.add_heading("企业运营分析报告", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_page_break()

    # 2. 自动目录（Heading样式可用于Word自动更新目录）
    doc.add_heading("目录", level=1)
    doc.add_paragraph(
        "1. 客户反馈分析\n2. 社交媒体情绪分析\n3. 智能总结", style='Normal'
    )
    doc.add_page_break()

    # 3. 客户反馈分析
    doc.add_heading("客户反馈分析", level=1)
    for para in feedback_text.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_page_break()

    # 3. 社交媒体情绪分析
    doc.add_heading("社交媒体情绪分析", level=1)
    for para in social_media_text.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_page_break()

    # 4. 图表图片（可选）
    if chart_path and os.path.exists(chart_path):
        doc.add_heading("附图：情绪分布图", level=2)
        doc.add_picture(chart_path, width=Inches(5.5))
        doc.add_page_break()

    # 5. LLM 总结段落
    doc.add_heading("智能总结", level=1)
    for para in llm_summary.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_page_break()

    # 保存文档
    output_path = os.path.join("static", "reports", filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path